import os
import uuid
import logging
from typing import List, Dict, Any, Tuple
import pypdf
import docx

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LC_Document
from app.core.config import settings

logger = logging.getLogger(__name__)

def get_chroma_db() -> Chroma:
    """Initialize or load the persistent Chroma DB vector store."""
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001",
        google_api_key=settings.GEMINI_API_KEY
    )
    return Chroma(
        persist_directory=settings.CHROMA_PERSIST_DIRECTORY,
        embedding_function=embeddings,
        collection_name="rag_documents"
    )

def extract_text_from_file(file_path: str, file_type: str) -> List[Tuple[str, int]]:
    """
    Extract text from a file based on its file extension.
    Returns a list of tuples containing (page_text, page_number).
    """
    pages_data: List[Tuple[str, int]] = []
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File tidak ditemukan di path: {file_path}")

    ext = file_type.lower()
    
    try:
        if ext == "pdf":
            reader = pypdf.PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_data.append((text, page_idx + 1))
                    
        elif ext in ["docx", "doc"]:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # DOCX has no fixed pagination in Python parsing, treat as 1 continuous page
            if full_text:
                pages_data.append(("\n".join(full_text), 1))
                
        elif ext in ["txt", "md", "csv"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                if content.strip():
                    pages_data.append((content, 1))
        else:
            raise ValueError(f"Tipe file '{ext}' tidak didukung.")
            
    except Exception as e:
        logger.error(f"Error parsing file {file_path}: {str(e)}")
        raise RuntimeError(f"Gagal mengekstrak teks dari file: {str(e)}")
        
    return pages_data

async def process_and_embed_document(
    user_id: uuid.UUID,
    document_id: uuid.UUID,
    file_path: str,
    file_type: str,
    filename: str
) -> bool:
    """
    RAG Pipeline:
    1. Extract text and split by pages.
    2. Split text into overlapping chunks using RecursiveCharacterTextSplitter.
    3. Generate embeddings and ingest chunks into ChromaDB with multi-tenant metadata.
    """
    try:
        logger.info(f"Memulai pemrosesan dokumen: {filename} (ID: {document_id})")
        
        # 1. Extract text
        pages = extract_text_from_file(file_path, file_type)
        if not pages:
            raise ValueError("File tidak memiliki konten teks yang dapat diekstrak.")
            
        # 2. Chunking
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        
        langchain_docs: List[LC_Document] = []
        
        for text, page_num in pages:
            chunks = splitter.split_text(text)
            for chunk_idx, chunk in enumerate(chunks):
                metadata = {
                    "user_id": str(user_id),
                    "document_id": str(document_id),
                    "filename": filename,
                    "page": page_num,
                    "chunk_index": chunk_idx
                }
                langchain_docs.append(LC_Document(page_content=chunk, metadata=metadata))
        
        if not langchain_docs:
            raise ValueError("Tidak ada chunk teks yang berhasil dihasilkan.")
            
        logger.info(f"Berhasil menghasilkan {len(langchain_docs)} chunks untuk {filename}")
        
        # 3. Save to Vector DB (Chroma)
        db = get_chroma_db()
        # Ingest docs
        db.add_documents(langchain_docs)
        logger.info(f"Ingest {filename} ke ChromaDB selesai sukses.")
        return True
        
    except Exception as e:
        logger.error(f"Gagal memproses embedding untuk dokumen {filename}: {str(e)}")
        return False

def delete_document_embeddings(document_id: uuid.UUID) -> None:
    """Delete document chunks from ChromaDB completely using document_id as filter."""
    try:
        db = get_chroma_db()
        # Chroma vector store exposes the ._collection to run custom queries/deletions
        db._collection.delete(where={"document_id": str(document_id)})
        logger.info(f"Embeddings untuk dokumen ID {document_id} berhasil dihapus dari ChromaDB.")
    except Exception as e:
        logger.error(f"Gagal menghapus embeddings dokumen {document_id} dari ChromaDB: {str(e)}")
        # Raise so parent handler can take actions if needed
        raise
